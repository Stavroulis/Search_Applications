import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from PIL import Image

import json
from datetime import date
from pathlib import Path
import pandas as pd

# Ensure filename is in session state BEFORE using it
if "filename" not in st.session_state:
    st.warning("No file selected. Please go to the main page.")
    st.stop()

filename = st.session_state["filename"]          
directory = Path(f"data/{filename}")
directory.mkdir(parents=True, exist_ok=True)  # Ensure directory exists
file_path = directory / f"Summary_{filename}.json"

# Load existing file if it exists
def load_json():
    if file_path.exists():
        with open(file_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def create_word_doc(filename, data):
    document = Document()
    
    # Setting page dimensions
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    
    # Add document metadata
    document.core_properties.author = "Dr. St^2"
    
    # Create a title with filename and date
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    filename_run = title_paragraph.add_run(filename)
    filename_run.font.name = "Arial"
    filename_run.bold = True
    filename_run.font.size = Pt(16)
    title_paragraph.add_run("\t" * 7)
    date_run = title_paragraph.add_run(f"{date.today()}")
    date_run.font.name = "Arial"
    date_run.bold = True
    date_run.font.size = Pt(16)
    
    # Create the table with 2 columns
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    
    # Define the labels for the left column and alternating row colors
    labels = [
        "Independent Claims", "Ptbs", "Solution", "Technical Effect", "Keywords", 
        "Classes", "Remarks", "Unity", "Prior Art"
    ]
    
    # Add alternating background colors and font formatting
    for i, label in enumerate(labels):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(data.get(label, ""))
        
        # Style left column (label cell)
        run_left = row[0].paragraphs[0].runs[0]
        run_left.font.name = "Arial"
        run_left.font.bold = True
        run_left.font.size = Pt(14)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Insert an empty line in the left column by adding a new paragraph
        row[0].add_paragraph()  # Empty line in the left column

        # Style right column (value cell)
        run_right = row[1].paragraphs[0].runs[0]
        run_right.font.name = "Arial"
        run_right.font.size = Pt(12)

        # Apply alternating background colors to the cells
        if i % 2 == 0:
            # Light blue color for even rows
            for cell in row:
                cell._element.get_or_add_tcPr().append(create_shading_element("D9EAF7"))  # Light blue color code
        else:
            # White color for odd rows
            for cell in row:
                cell._element.get_or_add_tcPr().append(create_shading_element("FFFFFF"))  # White color code

    # Check if the image exists in the same directory
    image_path = directory / f"appl_image_{filename}.png"
    
    # Add a new row with a single merged cell
    image_row = table.add_row().cells
    image_cell = image_row[0]
    image_cell.merge(image_row[1])  # Merge the two columns

    if image_path.is_file():
        # Open the image to get its original size
        with Image.open(image_path) as img:
            img_width, img_height = img.size  # Get original dimensions
            
            # Max dimensions for the document
            max_width_mm = 140
            max_height_mm = 100
            
            # Convert mm to pixels (assuming 96 DPI)
            mm_to_px = lambda mm: int((mm / 25.4) * 96)  # Convert mm to pixels
            
            max_width_px = mm_to_px(max_width_mm)
            max_height_px = mm_to_px(max_height_mm)
            
            # Calculate the new dimensions while keeping aspect ratio
            aspect_ratio = img_width / img_height
            
            if img_width > max_width_px or img_height > max_height_px:
                if img_width / max_width_px > img_height / max_height_px:
                    new_width = max_width_px
                    new_height = int(new_width / aspect_ratio)
                else:
                    new_height = max_height_px
                    new_width = int(new_height * aspect_ratio)
            else:
                new_width, new_height = img_width, img_height  # No scaling needed
            
            # Convert pixels back to mm
            new_width_mm = (new_width / 96) * 25.4
            new_height_mm = (new_height / 96) * 25.4
            
        # Add the image to the merged cell with adjusted dimensions
        paragraph = image_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Mm(new_width_mm), height=Mm(new_height_mm))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Add placeholder text
        image_cell.text = "You did not provide an application image."
        run = image_cell.paragraphs[0].runs[0]
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for emphasis
        image_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the Word document
    docx_filename = directory / f"Summary_{filename}.docx"
    
    # Debugging line to print out the path
    print(f"Saving the document at: {docx_filename}")
    
    # Save the Word document
    document.save(docx_filename)
    
    return docx_filename

def create_shading_element(color):
    """Creates a shading XML element for table cell background."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)  # Pass the color as a hex string
    return shading

data = load_json()

if not data:
    st.warning("No data found! Please check the file.")
    st.stop()

# Extract necessary data
extracted_data = {key: data.get(key, "") for key in [
    "Independent Claims", "Ptbs", "Technical Effect", "Solution", "Keywords", 
    "Classes", "Unity", "Remarks", "Prior Art", "Nr. Claims", "Date"]
}
extracted_data["Markers"] = data.get("Markers", {})
text_data = json.dumps(extracted_data, indent=4)

# Display JSON summary for debugging purposes
#st.markdown(f"<h3 style='color: red;'>FILE: {filename}</h3>", unsafe_allow_html=True)
#st.text_area("Summary", value=text_data, height=400)

ross_data = {key: data.get(key, "") for key in [
    "Ptbs", "Technical Effect", "Solution", "Keywords", "Classes", "Remarks"]
}

# Check if ross_data contains any non-empty values and set ross_text accordingly
ross_text = "\n".join(str(value) for value in ross_data.values())

# If no values in ross_text, set ross_text to None to trigger placeholder
if not ross_text.strip():
    ross_text = None

# Display only the required information in a text area with placeholder
st.text_area("RoSS Summary", value=ross_text, height=200, placeholder="If this is empty you must add information in the tab 'General'")

# Button to create the Word document
if st.button("Create Word"):
    try:
        # Create the document
        docx_path = create_word_doc(filename, extracted_data)
        if docx_path is not None:
            st.success("Word document created successfully!")
            st.session_state.docx_path = docx_path  # Save the path to session state for later use
        else:
            st.error("Failed to create the document.")
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")

# Button to download the Word document
if 'docx_path' in st.session_state:
    docx_path = st.session_state.docx_path
    with open(docx_path, "rb") as file:
        st.download_button(
            label="Download Word Document",
            data=file,
            file_name=f"Summary_{filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
